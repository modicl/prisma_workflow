"""
Cargador de documentos: convierte PDF, DOCX y JSON a texto plano
para inyectar en el session.state de Google ADK.

PDFs pasan por Gemini Files API (OCR multimodal) para manejar documentos
digitales, escaneados y mixtos. DOCX se extrae localmente con python-docx;
si el texto extraído es insuficiente (DOCX escaneado), se aplica OCR inline
sobre las imágenes embebidas usando Gemini.
"""

import base64
import json
import os
import zipfile
from pathlib import Path

from docx import Document
from google import genai
from google.genai import types as genai_types


_PDF_MIME = "application/pdf"
_MODEL = "gemini-2.5-flash-lite"
MIN_TEXT_CHARS = 200  # umbral mínimo para considerar que python-docx extrajo contenido útil

_IMAGE_MIMES = {
    ".png":  "image/png",
    ".jpg":  "image/jpeg",
    ".jpeg": "image/jpeg",
    ".gif":  "image/gif",
    ".webp": "image/webp",
    ".bmp":  "image/bmp",
}


def load_document(path: str, label: str | None = None) -> str:
    """
    Carga un documento y retorna su contenido como texto plano.

    Args:
        path: Ruta del archivo en disco (.pdf, .docx, .doc, .json).
        label: Etiqueta descriptiva para el encabezado del texto extraído.

    Returns:
        Texto extraído listo para inyectar en session.state de un LLM.

    Raises:
        FileNotFoundError: Si el archivo no existe.
        ValueError: Si la extensión no está soportada o no se pudo extraer texto.
    """
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"Archivo no encontrado: {path}")

    suffix = p.suffix.lower()

    if suffix == ".pdf":
        return _load_pdf_via_gemini(p, label)
    elif suffix == ".docx":
        return _load_docx(p, label)
    elif suffix == ".doc":
        raise ValueError(
            f"El archivo '{p.name}' está en formato .doc (Word 97-2003), que no es compatible.\n"
            "Por favor guárdalo como .docx (Archivo → Guardar como → Word (.docx)) o como .pdf e intenta de nuevo."
        )
    elif suffix == ".json":
        return _load_json(p, label)
    else:
        raise ValueError(f"Formato no soportado: '{suffix}'. Use .pdf, .docx o .json")


def _load_pdf_via_gemini(path: Path, label: str | None) -> str:
    """
    Sube el PDF a la Files API de Gemini y extrae su contenido completo.

    Aplica OCR multimodal que maneja correctamente PDFs digitales, escaneados
    y mixtos (texto + imágenes en el mismo documento). El archivo se elimina
    de la nube inmediatamente tras la extracción (requerimiento PII).
    """
    client = genai.Client(api_key=os.environ.get("GOOGLE_API_KEY"))

    print(f"  → Subiendo {label or path.name} a Google Files API...")

    uploaded = None
    try:
        with open(path, "rb") as f:
            uploaded = client.files.upload(
                file=f,
                config={"mime_type": _PDF_MIME, "display_name": path.name},
            )

        print(f"  → Extrayendo contenido con Gemini...  (puede tardar 20-60s según el documento)")
        response = client.models.generate_content(
            model=_MODEL,
            contents=[
                uploaded,
                "Extrae y transcribe el texto completo de este documento. "
                "Incluye párrafos, tablas, cuadros de texto, encabezados, pies de página "
                "y cualquier texto visible en imágenes o páginas escaneadas. "
                "Transcribe fielmente sin agregar interpretaciones ni resúmenes.",
            ],
            config=genai_types.GenerateContentConfig(
                thinking_config=genai_types.ThinkingConfig(thinking_budget=0),
            ),
        )
    finally:
        if uploaded is not None:
            print(f"  → Eliminando archivo de la nube...")
            try:
                client.files.delete(name=uploaded.name)
            except Exception:
                pass

    extracted = response.text
    if not extracted and response.candidates:
        content = response.candidates[0].content
        if content and content.parts:
            extracted = "".join(
                p.text for p in content.parts
                if hasattr(p, "text") and p.text and not getattr(p, "thought", False)
            )

    if not extracted:
        finish_reason = None
        if response.candidates:
            finish_reason = getattr(response.candidates[0], "finish_reason", None)
        raise ValueError(
            f"Gemini no pudo extraer texto de '{path.name}' "
            f"(finish_reason={finish_reason}). "
            "El documento puede estar protegido con contraseña o dañado."
        )

    prefix = f"[Documento PDF: {label or path.name}]\n\n" if label else f"[{path.name}]\n\n"
    text = prefix + extracted
    print(f"  ✓ {label or path.name} cargado ({len(text):,} caracteres) [Gemini]")
    return text


def _load_docx(path: Path, label: str | None) -> str:
    """
    Extrae texto de un DOCX. Usa python-docx como método principal.
    Si el texto extraído es insuficiente (DOCX escaneado), aplica OCR
    sobre las imágenes embebidas usando Gemini inline.
    """
    doc = Document(str(path))

    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    text_from_xml = "\n".join(parts)

    if len(text_from_xml) >= MIN_TEXT_CHARS:
        prefix = f"[Documento DOCX: {label or path.name}]\n\n" if label else f"[{path.name}]\n\n"
        text = prefix + text_from_xml
        print(f"  ✓ {label or path.name} cargado ({len(text):,} caracteres) [local]")
        return text

    # Texto insuficiente — intentar OCR sobre imágenes embebidas
    print(f"  → Texto insuficiente en '{path.name}' ({len(text_from_xml)} chars). "
          "Buscando imágenes para OCR...")
    images = _extract_docx_images(path)

    if not images:
        if text_from_xml.strip():
            prefix = f"[Documento DOCX: {label or path.name}]\n\n" if label else f"[{path.name}]\n\n"
            return prefix + text_from_xml
        raise ValueError(
            f"No se pudo extraer texto de '{path.name}'. "
            "El archivo parece estar vacío. "
            "Si contiene páginas escaneadas, conviértelo a PDF e intenta de nuevo."
        )

    ocr_text = _ocr_images_with_gemini(images, path.name)
    combined = "\n\n".join(filter(None, [text_from_xml, ocr_text]))

    prefix = f"[Documento DOCX: {label or path.name}]\n\n" if label else f"[{path.name}]\n\n"
    text = prefix + combined
    print(f"  ✓ {label or path.name} cargado ({len(text):,} caracteres) [local + OCR Gemini]")
    return text


def _extract_docx_images(path: Path) -> list[dict]:
    """
    Extrae imágenes embebidas de un DOCX (que internamente es un ZIP).
    Retorna lista de dicts con 'mime_type' y 'data' (bytes).
    """
    images = []
    with zipfile.ZipFile(path, "r") as z:
        for name in z.namelist():
            if not name.startswith("word/media/"):
                continue
            ext = Path(name).suffix.lower()
            mime = _IMAGE_MIMES.get(ext)
            if mime is None:
                continue
            images.append({"mime_type": mime, "data": z.read(name)})
    return images


def _ocr_images_with_gemini(images: list[dict], filename: str) -> str:
    """
    Envía imágenes a Gemini como datos inline (sin Files API) y retorna el texto extraído.
    """
    client = genai.Client(api_key=os.environ.get("GOOGLE_API_KEY"))
    print(f"  → Aplicando OCR con Gemini sobre {len(images)} imagen(es) de '{filename}'...")

    ocr_parts: list = []
    for img in images:
        ocr_parts.append({
            "inline_data": {
                "mime_type": img["mime_type"],
                "data": base64.b64encode(img["data"]).decode("utf-8"),
            }
        })
    ocr_parts.append(
        "Extrae y transcribe todo el texto visible en estas imágenes. "
        "Incluye párrafos, tablas, encabezados y cualquier texto legible. "
        "Transcribe fielmente sin agregar interpretaciones ni resúmenes."
    )

    response = client.models.generate_content(
        model=_MODEL,
        contents=ocr_parts,
        config=genai_types.GenerateContentConfig(
            thinking_config=genai_types.ThinkingConfig(thinking_budget=0),
        ),
    )

    extracted = response.text
    if not extracted and response.candidates:
        content = response.candidates[0].content
        if content and content.parts:
            extracted = "".join(
                p.text for p in content.parts
                if hasattr(p, "text") and p.text and not getattr(p, "thought", False)
            )
    return extracted or ""


def _load_json(path: Path, label: str | None) -> str:
    """Carga un PACI estructurado en JSON y lo formatea para el LLM."""
    with open(path, "r", encoding="utf-8") as f:
        data = json.load(f)

    formatted = json.dumps(data, ensure_ascii=False, indent=2)
    prefix = f"[PACI en formato JSON — {label or path.name}]\n\n"
    return prefix + formatted
