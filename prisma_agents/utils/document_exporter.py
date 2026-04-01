"""
Exportador de documentos: guarda los resultados del flujo en un archivo de Word (.docx) 
con formato limpio, eliminando el markdown bruto (##, **, tablas) y sin incluir el perfil.
"""
import os
import re
from docx import Document
from docx.shared import Pt

def _add_formatted_runs(paragraph, text):
    """
    Toma un texto con markdown de negrita (**texto**) y lo añade 
    al párrafo de python-docx con los runs correspondientes.
    """
    # Expresión regular para encontrar texto entre ** **
    parts = re.split(r'\*\*(.*?)\*\*', text)
    for i, part in enumerate(parts):
        if part:
            run = paragraph.add_run(part)
            if i % 2 == 1:
                run.bold = True

def _add_markdown_to_doc(doc, markdown_text):
    """
    Parsea markdown básico (Títulos, Listas, Tablas, Negritas) y 
    lo agrega a un documento de python-docx limpio.
    """
    lines = markdown_text.split('\n')
    in_table = False
    table_rows = []

    def flush_table():
        if not table_rows:
            return
        # Crear la tabla
        table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        table.style = 'Table Grid' # Agrega los bordes a la tabla
        
        for row_idx, row_data in enumerate(table_rows):
            for col_idx, cell_data in enumerate(row_data):
                if col_idx >= len(table.columns):
                    continue
                cell = table.cell(row_idx, col_idx)
                p = cell.paragraphs[0]
                _add_formatted_runs(p, cell_data)
                
                # Poner la primera fila (encabezado) en negrita
                if row_idx == 0:
                    for run in p.runs:
                        run.bold = True
        table_rows.clear()

    for line in lines:
        stripped = line.strip()
        
        # Ignorar líneas separadoras
        if stripped == '---':
            continue

        # Detectar tabla
        if stripped.startswith('|') and stripped.endswith('|'):
            # Ignorar la línea divisoria del encabezado de la tabla (ej. |:---|:---|)
            if re.search(r'\|[\s\-:]+\|', stripped):
                continue
            
            # Extraer columnas eliminando los |
            columns = [c.strip() for c in stripped.split('|')[1:-1]]
            table_rows.append(columns)
            in_table = True
            continue
        else:
            if in_table:
                flush_table()
                in_table = False

        if not stripped:
            continue

        # Calcular nivel de indentación para sub-listas
        indent_level = len(line) - len(line.lstrip())

        # Procesar encabezados (###)
        if stripped.startswith('#'):
            level = len(stripped.split(' ')[0])
            texto_limpio = stripped.lstrip('#').strip()
            # python-docx soporta heading levels 0 a 9
            p = doc.add_heading('', level=min(level, 9))
            _add_formatted_runs(p, texto_limpio)

        # Procesar viñetas (* o -)
        elif stripped.startswith('* ') or stripped.startswith('- '):
            texto_limpio = stripped[2:].strip()
            
            # Si hay 4 o más espacios de indentación, usamos viñeta nivel 2
            style = 'List Bullet 2' if indent_level >= 4 else 'List Bullet'
            try:
                p = doc.add_paragraph(style=style)
            except KeyError:
                p = doc.add_paragraph(style='List Bullet') # Fallback seguro
                
            _add_formatted_runs(p, texto_limpio)
            
        else:
            # Párrafo normal
            p = doc.add_paragraph()
            _add_formatted_runs(p, stripped)

    # Si el documento terminaba en una tabla, asegurarnos de volcarla
    if in_table:
        flush_table()


def export_results_to_docx(results: dict, output_filename: str = "resultado_paci.docx") -> str:
    """
    Crea un nuevo documento Word con la Planificación Adaptada y la Rúbrica Final,
    eliminando el markdown bruto (parseando a formato MS Word nativo) y omitiendo el perfil PACI.
    """
    doc = Document()
    
    # Título principal
    doc.add_heading('Documento de Apoyo Docente - Material Adaptado', 0)

    # 1. Planificación Adaptada
    doc.add_heading('1. Planificación de la Actividad Adaptada', level=1)
    if results.get("planificacion_adaptada"):
        _add_markdown_to_doc(doc, results["planificacion_adaptada"])
    else:
        doc.add_paragraph("(No se generó una planificación adaptada)")

    doc.add_page_break() # Añadir un salto de página entre secciones

    # 2. Rúbrica Final
    doc.add_heading('2. Rúbrica Final de Evaluación', level=1)
    if results.get("rubrica_final"):
        _add_markdown_to_doc(doc, results["rubrica_final"])
    else:
        doc.add_paragraph("(No se generó rúbrica)")

    # Se omitió el Perfil PACI para no enviar información redundante al docente
    
    doc.save(output_filename)
    return os.path.abspath(output_filename)
