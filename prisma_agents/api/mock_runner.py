"""
Mock workflow runner for dev/UX testing — no LLM calls, no S3, no tokens consumed.

Activate by setting school_id to one of:
  __mock_success__   → Full flow: A1 → A2 → HITL → A3 → Crítico → descarga
  __mock_fast__      → No HITL, completes in ~6s
  __mock_degraded__  → Completes as degraded (best effort)
  __mock_error__     → Fails with timeout error
"""

import asyncio
import os
import tempfile

from docx import Document

from api.session_store import sync_to_dynamo

STEP_DELAY = 1.2


def _push(session_data, content, role="system"):
    msg = {"role": role, "content": content}
    session_data.messages.append(msg)
    session_data.event_queue.put_nowait({"type": "message", **msg})


def _make_placeholder_docx() -> str:
    doc = Document()
    doc.add_heading("Rúbrica Adaptada — Simulación de Prueba", 0)
    doc.add_paragraph(
        "Este documento fue generado por el simulador de desarrollo de P.R.I.S.M.A. "
        "No representa una rúbrica real."
    )
    doc.add_heading("Criterios de Evaluación (simulados)", level=1)
    for i in range(1, 4):
        doc.add_paragraph(f"Criterio {i}: Descripción simulada del criterio de evaluación.", style="List Bullet")
    fd, path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    doc.save(path)
    return path


async def run_mock_workflow(session_id: str, session_data, school_id: str) -> None:
    scenario = school_id.replace("__mock_", "").rstrip("_")
    dispatch = {
        "success":  _mock_success,
        "fast":     _mock_fast,
        "degraded": _mock_degraded,
        "error":    _mock_error,
    }
    await dispatch.get(scenario, _mock_success)(session_id, session_data)


# ── Scenarios ─────────────────────────────────────────────────────────────────

async def _mock_success(session_id, session_data):
    """A1 → A2 → HITL (espera aprobación docente) → A3 → Crítico → completado."""
    _push(session_data, "Documentos recibidos. Iniciando análisis del PACI...")
    sync_to_dynamo(session_id, session_data)
    await asyncio.sleep(STEP_DELAY)

    for agent, msg in [
        ("Agente 1", "Analizando PACI..."),
        ("Agente 2", "Adaptando material educativo..."),
    ]:
        session_data.event_queue.put_nowait({"type": "agent_start", "agent": agent, "message": msg})
        await asyncio.sleep(STEP_DELAY * 2)
        session_data.event_queue.put_nowait({"type": "agent_end", "agent": agent})

    hitl_data = {
        "perfil_paci": (
            "[SIMULADO] Estudiante con TEA grado 2, 4° básico. "
            "Requiere apoyos visuales, instrucciones cortas y estructura predecible. "
            "Fortalezas en memoria visual y rutinas."
        ),
        "planificacion_adaptada": (
            "[SIMULADO] Unidad: Fracciones. Material adaptado con pictogramas, "
            "pasos secuenciados y ejemplos concretos con material manipulativo. "
            "Reducción de ítems por evaluación."
        ),
        "attempt": 1,
        "max_attempts": 3,
    }
    session_data.hitl_data = hitl_data
    session_data.phase = "awaiting_hitl"
    _push(session_data, "Revisión requerida — intento 1 de 3. Por favor revise el análisis y la planificación.")
    session_data.event_queue.put_nowait({
        "type": "hitl_required",
        "attempt": 1,
        "max_attempts": 3,
        "hitl_data": hitl_data,
    })
    sync_to_dynamo(session_id, session_data)

    response = await session_data.hitl_response_queue.get()
    session_data.phase = "running"
    session_data.hitl_data = None
    sync_to_dynamo(session_id, session_data)

    if not response.get("approved", False):
        session_data.phase = "error"
        session_data.workflow_status = "hitl_rejected"
        session_data.event_queue.put_nowait({"type": "error", "message": "Proceso cancelado por el docente."})
        sync_to_dynamo(session_id, session_data)
        return

    for agent, msg in [
        ("Agente 3", "Generando rúbrica adaptada..."),
        ("Agente Crítico", "Evaluando calidad de la rúbrica..."),
    ]:
        session_data.event_queue.put_nowait({"type": "agent_start", "agent": agent, "message": msg})
        await asyncio.sleep(STEP_DELAY * 2)
        session_data.event_queue.put_nowait({"type": "agent_end", "agent": agent})

    _push(session_data, "✅ Proceso completado. La rúbrica adaptada está lista para descargar.", role="agent")
    session_data.docx_path = _make_placeholder_docx()
    session_data.phase = "completed"
    session_data.workflow_status = "success"
    session_data.event_queue.put_nowait({"type": "completed", "workflow_status": "success"})
    sync_to_dynamo(session_id, session_data)


async def _mock_fast(session_id, session_data):
    """Sin HITL. Completa en ~6 segundos."""
    _push(session_data, "Documentos recibidos. Iniciando análisis del PACI...")
    sync_to_dynamo(session_id, session_data)
    await asyncio.sleep(STEP_DELAY)

    for agent, msg in [
        ("Agente 1", "Analizando PACI..."),
        ("Agente 2", "Adaptando material educativo..."),
        ("Agente 3", "Generando rúbrica adaptada..."),
        ("Agente Crítico", "Evaluando calidad de la rúbrica..."),
    ]:
        session_data.event_queue.put_nowait({"type": "agent_start", "agent": agent, "message": msg})
        await asyncio.sleep(STEP_DELAY)
        session_data.event_queue.put_nowait({"type": "agent_end", "agent": agent})

    _push(session_data, "✅ Proceso completado. La rúbrica adaptada está lista para descargar.", role="agent")
    session_data.docx_path = _make_placeholder_docx()
    session_data.phase = "completed"
    session_data.workflow_status = "success"
    session_data.event_queue.put_nowait({"type": "completed", "workflow_status": "success"})
    sync_to_dynamo(session_id, session_data)


async def _mock_degraded(session_id, session_data):
    """Completa como mejor esfuerzo (degraded)."""
    _push(session_data, "Documentos recibidos. Iniciando análisis del PACI...")
    sync_to_dynamo(session_id, session_data)
    await asyncio.sleep(STEP_DELAY)

    for agent, msg in [
        ("Agente 1", "Analizando PACI..."),
        ("Agente 2", "Adaptando material educativo..."),
        ("Agente 3", "Generando rúbrica..."),
        ("Agente Crítico", "Evaluando calidad... (criterios no cumplidos, usando mejor esfuerzo)"),
    ]:
        session_data.event_queue.put_nowait({"type": "agent_start", "agent": agent, "message": msg})
        await asyncio.sleep(STEP_DELAY)
        session_data.event_queue.put_nowait({"type": "agent_end", "agent": agent})

    _push(
        session_data,
        "⚠️ Proceso completado. La rúbrica fue generada como mejor esfuerzo y no superó "
        "todos los criterios de calidad. Revise el documento antes de usarlo.",
        role="agent",
    )
    session_data.docx_path = _make_placeholder_docx()
    session_data.phase = "completed"
    session_data.workflow_status = "degraded"
    session_data.event_queue.put_nowait({"type": "completed", "workflow_status": "degraded"})
    sync_to_dynamo(session_id, session_data)


async def _mock_error(session_id, session_data):
    """El Agente 1 falla por timeout."""
    _push(session_data, "Documentos recibidos. Iniciando análisis del PACI...")
    sync_to_dynamo(session_id, session_data)
    await asyncio.sleep(STEP_DELAY)

    session_data.event_queue.put_nowait({"type": "agent_start", "agent": "Agente 1", "message": "Analizando PACI..."})
    await asyncio.sleep(STEP_DELAY * 2.5)

    error_msg = "El servicio de IA no respondió a tiempo. Intente nuevamente."
    session_data.phase = "error"
    session_data.workflow_status = "error"
    session_data.error = error_msg
    _push(session_data, "❌ El procesamiento fue interrumpido por un error del servidor.")
    session_data.event_queue.put_nowait({"type": "error", "message": error_msg})
    sync_to_dynamo(session_id, session_data)
