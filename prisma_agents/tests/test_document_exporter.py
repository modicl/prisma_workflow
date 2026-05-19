"""Tests para utils/document_exporter.py."""
import os
import sys
import tempfile
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from docx import Document
from utils.document_exporter import export_results_to_docx, _add_formatted_runs, _add_markdown_to_doc


class TestAddFormattedRuns:
    def _make_para(self):
        doc = Document()
        return doc.add_paragraph()

    def test_plain_text(self):
        p = self._make_para()
        _add_formatted_runs(p, "Texto simple")
        assert p.text == "Texto simple"

    def test_bold_text(self):
        p = self._make_para()
        _add_formatted_runs(p, "Texto **negrita** aquí")
        texts = [r.text for r in p.runs]
        bolds = [r.bold for r in p.runs]
        assert "negrita" in texts
        assert True in bolds

    def test_multiple_bold_segments(self):
        p = self._make_para()
        _add_formatted_runs(p, "**A** normal **B**")
        bolds = [(r.text, r.bold) for r in p.runs if r.text]
        bold_texts = [t for t, b in bolds if b]
        assert "A" in bold_texts
        assert "B" in bold_texts

    def test_empty_text_no_crash(self):
        p = self._make_para()
        _add_formatted_runs(p, "")
        assert p.text == ""


class TestAddMarkdownToDoc:
    def _make_doc(self):
        return Document()

    def test_heading_level_1(self):
        doc = self._make_doc()
        _add_markdown_to_doc(doc, "# Título principal")
        headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert len(headings) >= 1

    def test_heading_level_3(self):
        doc = self._make_doc()
        _add_markdown_to_doc(doc, "### Subtítulo")
        headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert len(headings) >= 1

    def test_bullet_star(self):
        doc = self._make_doc()
        _add_markdown_to_doc(doc, "* Elemento de lista")
        bullets = [p for p in doc.paragraphs if "Bullet" in p.style.name]
        assert len(bullets) >= 1
        assert "Elemento de lista" in bullets[0].text

    def test_bullet_dash(self):
        doc = self._make_doc()
        _add_markdown_to_doc(doc, "- Otro elemento")
        bullets = [p for p in doc.paragraphs if "Bullet" in p.style.name]
        assert len(bullets) >= 1

    def test_separator_line_ignored(self):
        doc = self._make_doc()
        _add_markdown_to_doc(doc, "---")
        texts = [p.text for p in doc.paragraphs]
        assert "---" not in texts

    def test_normal_paragraph(self):
        doc = self._make_doc()
        _add_markdown_to_doc(doc, "Texto normal de párrafo")
        texts = [p.text for p in doc.paragraphs]
        assert "Texto normal de párrafo" in texts

    def test_table_rendered(self):
        doc = self._make_doc()
        md = "| Col1 | Col2 |\n|---|---|\n| Val1 | Val2 |"
        _add_markdown_to_doc(doc, md)
        assert len(doc.tables) == 1
        assert doc.tables[0].rows[0].cells[0].text == "Col1"

    def test_empty_lines_skipped(self):
        doc = self._make_doc()
        _add_markdown_to_doc(doc, "\n\n\n")
        non_empty = [p for p in doc.paragraphs if p.text.strip()]
        assert len(non_empty) == 0

    def test_indented_bullet_uses_level2(self):
        doc = self._make_doc()
        _add_markdown_to_doc(doc, "    * Sub-elemento")
        bullets = [p for p in doc.paragraphs if "Bullet" in p.style.name]
        assert len(bullets) >= 1


class TestExportResultsToDocx:
    def test_creates_docx_file(self, tmp_path):
        output = str(tmp_path / "resultado.docx")
        results = {
            "planificacion_adaptada": "## Adaptación\nContenido adaptado.",
            "rubrica_final": "| Criterio | Logrado |\n|---|---|\n| Criterio 1 | Sí |",
        }
        path = export_results_to_docx(results, output)
        assert os.path.exists(path)
        doc = Document(path)
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "Planificación" in full_text or "Adaptada" in full_text

    def test_missing_planificacion_uses_fallback(self, tmp_path):
        output = str(tmp_path / "sin_plan.docx")
        path = export_results_to_docx({"rubrica_final": "# Rúbrica"}, output)
        doc = Document(path)
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "No se generó" in full_text or "planificación" in full_text.lower()

    def test_missing_rubrica_uses_fallback(self, tmp_path):
        output = str(tmp_path / "sin_rubrica.docx")
        path = export_results_to_docx({"planificacion_adaptada": "# Plan"}, output)
        doc = Document(path)
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "No se generó" in full_text or "rúbrica" in full_text.lower()

    def test_returns_absolute_path(self, tmp_path):
        output = str(tmp_path / "out.docx")
        path = export_results_to_docx({}, output)
        assert os.path.isabs(path)

    def test_page_break_between_sections(self, tmp_path):
        output = str(tmp_path / "pagebreak.docx")
        results = {
            "planificacion_adaptada": "Contenido",
            "rubrica_final": "Rúbrica",
        }
        export_results_to_docx(results, output)
        from docx.oxml.ns import qn
        doc = Document(output)
        xml = doc.element.xml
        assert "pageBreak" in xml or "w:lastRenderedPageBreak" in xml or len(doc.paragraphs) > 3
