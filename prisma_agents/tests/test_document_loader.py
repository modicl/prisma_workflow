# prisma_agents/tests/test_document_loader.py
import io
import json
import os
import sys
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

import pytest
from unittest.mock import patch, MagicMock
from pathlib import Path

import utils.document_loader as dl


# ---------------------------------------------------------------------------
# Helpers compartidos
# ---------------------------------------------------------------------------

def _make_gemini_response(text: str) -> MagicMock:
    r = MagicMock()
    r.text = text
    r.candidates = []
    return r


def _make_gemini_client(response_text: str = "", side_effect: Exception | None = None) -> MagicMock:
    uploaded = MagicMock()
    uploaded.name = "files/mock123"
    client = MagicMock()
    client.files.upload.return_value = uploaded
    if side_effect is not None:
        client.models.generate_content.side_effect = side_effect
    else:
        client.models.generate_content.return_value = _make_gemini_response(response_text)
    return client


# ---------------------------------------------------------------------------
# load_document — enrutador
# ---------------------------------------------------------------------------

def test_load_document_raises_if_file_missing():
    with pytest.raises(FileNotFoundError):
        dl.load_document("/ruta/que/no/existe.pdf")


def test_load_document_raises_on_unsupported_extension(tmp_path):
    f = tmp_path / "doc.xls"
    f.write_bytes(b"data")
    with pytest.raises(ValueError, match="Formato no soportado"):
        dl.load_document(str(f))


# ---------------------------------------------------------------------------
# _load_pdf — ahora siempre usa Gemini (sin pdfplumber)
# ---------------------------------------------------------------------------

def test_load_pdf_uses_gemini_for_digital_pdf(tmp_path):
    """Un PDF con texto seleccionable debe pasar por Gemini, no pdfplumber."""
    pdf = tmp_path / "digital.pdf"
    pdf.write_bytes(b"%PDF-1.4 texto digital aqui")

    mock_client = _make_gemini_client("Texto extraído del PDF digital.")

    with patch("utils.document_loader.genai.Client", return_value=mock_client), \
         patch.dict(os.environ, {"GOOGLE_API_KEY": "test-key"}):
        result = dl.load_document(str(pdf), label="PACI Test")

    assert "Texto extraído del PDF digital." in result
    assert "PACI Test" in result
    # Gemini Files API debe haberse usado
    mock_client.files.upload.assert_called_once()
    # Archivo de Gemini debe eliminarse al finalizar (requerimiento PII)
    mock_client.files.delete.assert_called_once_with(name="files/mock123")


def test_load_pdf_compacts_extracted_text(tmp_path):
    """El texto extraído del PDF se compacta (whitespace/marcadores) antes de retornar."""
    pdf = tmp_path / "ruidoso.pdf"
    pdf.write_bytes(b"%PDF-1.4 texto")

    noisy = "Diagnóstico   TEA.\n\n\n\n3\nAdecuación curricular."
    mock_client = _make_gemini_client(noisy)

    with patch("utils.document_loader.genai.Client", return_value=mock_client), \
         patch.dict(os.environ, {"GOOGLE_API_KEY": "test-key"}):
        result = dl.load_document(str(pdf), label="PACI Test")

    # runs de espacios colapsados, blank lines colapsadas, marcador de página "3" eliminado
    assert "Diagnóstico TEA.\n\nAdecuación curricular." in result
    assert "\n\n\n" not in result


def test_load_pdf_uses_gemini_for_scanned_pdf(tmp_path):
    """Un PDF puramente escaneado (sin texto seleccionable) también pasa por Gemini."""
    pdf = tmp_path / "scanned.pdf"
    pdf.write_bytes(b"%PDF-1.4 binary image data")

    mock_client = _make_gemini_client("Texto OCR del escaneado.")

    with patch("utils.document_loader.genai.Client", return_value=mock_client), \
         patch.dict(os.environ, {"GOOGLE_API_KEY": "test-key"}):
        result = dl.load_document(str(pdf), label="Material")

    assert "Texto OCR del escaneado." in result
    assert "Material" in result


def test_load_pdf_deletes_gemini_file_even_on_generate_error(tmp_path):
    """El archivo en Gemini se elimina aunque generate_content falle (requerimiento PII)."""
    pdf = tmp_path / "bad.pdf"
    pdf.write_bytes(b"%PDF-1.4")

    uploaded = MagicMock()
    uploaded.name = "files/will_be_deleted"
    client = MagicMock()
    client.files.upload.return_value = uploaded
    client.models.generate_content.side_effect = RuntimeError("API error")

    with patch("utils.document_loader.genai.Client", return_value=client), \
         patch.dict(os.environ, {"GOOGLE_API_KEY": "test-key"}):
        with pytest.raises(RuntimeError):
            dl.load_document(str(pdf))

    client.files.delete.assert_called_once_with(name="files/will_be_deleted")


def test_load_pdf_raises_if_gemini_returns_empty(tmp_path):
    """Si Gemini retorna texto vacío, se lanza ValueError con mensaje claro."""
    pdf = tmp_path / "empty.pdf"
    pdf.write_bytes(b"%PDF-1.4")

    client = MagicMock()
    uploaded = MagicMock()
    uploaded.name = "files/empty"
    client.files.upload.return_value = uploaded
    empty_response = MagicMock()
    empty_response.text = ""
    empty_response.candidates = []
    client.models.generate_content.return_value = empty_response

    with patch("utils.document_loader.genai.Client", return_value=client), \
         patch.dict(os.environ, {"GOOGLE_API_KEY": "test-key"}):
        with pytest.raises(ValueError, match="no pudo extraer texto"):
            dl.load_document(str(pdf))


def test_load_pdf_no_pdfplumber_import(tmp_path):
    """pdfplumber no debe importarse en ningún punto del flujo de PDFs."""
    pdf = tmp_path / "test.pdf"
    pdf.write_bytes(b"%PDF-1.4")

    mock_client = _make_gemini_client("Texto de prueba.")

    import sys
    original_modules = dict(sys.modules)
    # Eliminamos pdfplumber del sys.modules para verificar que no se importa
    sys.modules.pop("pdfplumber", None)

    try:
        with patch("utils.document_loader.genai.Client", return_value=mock_client), \
             patch.dict(os.environ, {"GOOGLE_API_KEY": "test-key"}):
            dl.load_document(str(pdf))
        # Si pdfplumber fue importado, estará en sys.modules
        assert "pdfplumber" not in sys.modules, "pdfplumber fue importado — no debe usarse"
    finally:
        sys.modules.update(original_modules)


# ---------------------------------------------------------------------------
# _load_docx — usa python-docx local; Gemini inline OCR solo cuando hay imágenes
# ---------------------------------------------------------------------------

def _make_minimal_docx(tmp_path, filename="doc.docx", text="x") -> str:
    """Crea un DOCX válido usando python-docx para que python-docx pueda abrirlo."""
    from docx import Document as DocxDocument
    d = DocxDocument()
    if text:
        d.add_paragraph(text)
    buf = io.BytesIO()
    d.save(buf)
    path = tmp_path / filename
    path.write_bytes(buf.getvalue())
    return str(path)


def test_load_docx_uses_gemini_ocr_when_text_insufficient(tmp_path):
    """Si el texto local es < MIN_TEXT_CHARS y hay imágenes, usa Gemini OCR inline."""
    docx_path = _make_minimal_docx(tmp_path, text="Poco texto")  # < 200 chars

    fake_images = [{"mime_type": "image/png", "data": b"\x89PNG fake"}]
    mock_client = MagicMock()
    mock_client.models.generate_content.return_value.text = "OCR extraído desde imagen"

    with patch("utils.document_loader._extract_docx_images", return_value=fake_images), \
         patch("utils.document_loader.genai.Client", return_value=mock_client), \
         patch.dict(os.environ, {"GOOGLE_API_KEY": "test-key"}):
        result = dl.load_document(docx_path, label="Material DOCX")

    assert "OCR extraído desde imagen" in result
    mock_client.models.generate_content.assert_called_once()


def test_load_docx_gemini_ocr_error_propagates(tmp_path):
    """Si Gemini falla durante el OCR de imágenes, el error se propaga."""
    docx_path = _make_minimal_docx(tmp_path, text="x")  # < 200 chars

    fake_images = [{"mime_type": "image/png", "data": b"\x89PNG fake"}]
    mock_client = MagicMock()
    mock_client.models.generate_content.side_effect = RuntimeError("fallo Gemini OCR")

    with patch("utils.document_loader._extract_docx_images", return_value=fake_images), \
         patch("utils.document_loader.genai.Client", return_value=mock_client), \
         patch.dict(os.environ, {"GOOGLE_API_KEY": "test-key"}):
        with pytest.raises(RuntimeError, match="fallo Gemini OCR"):
            dl.load_document(docx_path)


def test_load_docx_with_sufficient_text_returns_local(tmp_path):
    """DOCX con texto suficiente (≥ 200 chars) se extrae localmente sin Gemini."""
    long_text = "Contenido educativo de prueba. " * 10  # > 200 chars
    docx_path = _make_minimal_docx(tmp_path, text=long_text)
    result = dl.load_document(docx_path, label="Guía de estudio")
    assert "Guía de estudio" in result
    assert "Contenido educativo de prueba." in result


def test_load_docx_no_text_no_images_raises_or_returns(tmp_path):
    """DOCX sin texto significativo y sin imágenes retorna el texto escaso o lanza error."""
    docx_path = _make_minimal_docx(tmp_path, text="")  # empty DOCX

    with patch("utils.document_loader._extract_docx_images", return_value=[]):
        try:
            result = dl.load_document(docx_path)
            # If it doesn't raise, it should return empty-ish content
            assert isinstance(result, str)
        except ValueError:
            pass  # acceptable — DOCX has no content


def test_load_docx_sparse_text_no_images_returns_sparse(tmp_path):
    """DOCX con texto escaso (< 200 chars) pero sin imágenes devuelve lo poco que hay."""
    short_text = "Breve."
    docx_path = _make_minimal_docx(tmp_path, text=short_text)

    with patch("utils.document_loader._extract_docx_images", return_value=[]):
        result = dl.load_document(docx_path)
    # Text is sparse but returned
    assert "Breve." in result


def test_extract_docx_images_no_media(tmp_path):
    """DOCX sin imágenes devuelve lista vacía."""
    from utils.document_loader import _extract_docx_images
    from pathlib import Path
    docx_path = Path(_make_minimal_docx(tmp_path, text="Texto"))
    images = _extract_docx_images(docx_path)
    assert images == []


def test_ocr_images_with_gemini_returns_text():
    """_ocr_images_with_gemini llama a Gemini con imágenes inline y retorna texto."""
    from utils.document_loader import _ocr_images_with_gemini
    mock_client = MagicMock()
    mock_client.models.generate_content.return_value.text = "Texto OCR"
    mock_client.models.generate_content.return_value.candidates = []

    with patch("utils.document_loader.genai.Client", return_value=mock_client), \
         patch.dict(os.environ, {"GOOGLE_API_KEY": "test-key"}):
        result = _ocr_images_with_gemini(
            [{"mime_type": "image/png", "data": b"\x89PNG"}], "doc.docx"
        )
    assert result == "Texto OCR"


def test_ocr_images_with_gemini_empty_text_fallback():
    """Si text está vacío, extrae de candidates."""
    from utils.document_loader import _ocr_images_with_gemini
    mock_response = MagicMock()
    mock_response.text = ""
    part = MagicMock()
    part.text = "De candidate"
    part.thought = False
    mock_response.candidates[0].content.parts = [part]

    mock_client = MagicMock()
    mock_client.models.generate_content.return_value = mock_response

    with patch("utils.document_loader.genai.Client", return_value=mock_client), \
         patch.dict(os.environ, {"GOOGLE_API_KEY": "test-key"}):
        result = _ocr_images_with_gemini(
            [{"mime_type": "image/png", "data": b"\x89PNG"}], "doc.docx"
        )
    assert "De candidate" in result


# ---------------------------------------------------------------------------
# .doc — sin cambios (sigue siendo ValueError)
# ---------------------------------------------------------------------------

def test_load_pdf_empty_text_raises_value_error(tmp_path):
    """Cubre lines 111-118, 126: Gemini devuelve texto vacío → ValueError."""
    pdf = tmp_path / "empty.pdf"
    pdf.write_bytes(b"%PDF-1.4 fake")

    mock_response = MagicMock()
    mock_response.text = ""
    mock_response.candidates = []

    uploaded = MagicMock()
    uploaded.name = "files/empty123"
    mock_client = _make_gemini_client("")
    mock_client.files.upload.return_value = uploaded
    mock_client.models.generate_content.return_value = mock_response

    with patch("utils.document_loader.genai.Client", return_value=mock_client), \
         patch.dict(os.environ, {"GOOGLE_API_KEY": "test-key"}):
        with pytest.raises(ValueError, match="no pudo extraer texto"):
            dl.load_document(str(pdf))


def test_load_docx_with_table_extracts_table_text(tmp_path):
    """Cubre lines 152-155: extracción de tabla en DOCX local."""
    from docx import Document as DocxDocument
    d = DocxDocument()
    table = d.add_table(rows=2, cols=2)
    long_cell = "Texto de celda educativa para tabla. " * 5  # contributes to >= 200 chars
    table.cell(0, 0).text = long_cell
    table.cell(0, 1).text = "Columna B"
    table.cell(1, 0).text = "Fila 2 Col A"
    table.cell(1, 1).text = "Fila 2 Col B"
    buf = io.BytesIO()
    d.save(buf)
    path = tmp_path / "con_tabla.docx"
    path.write_bytes(buf.getvalue())

    result = dl.load_document(str(path), label="Material con tabla")
    assert "Columna B" in result or long_cell[:20] in result


def test_extract_docx_images_with_image(tmp_path):
    """Cubre lines 199-203: DOCX con imagen embebida en word/media/."""
    import zipfile as _zf
    from utils.document_loader import _extract_docx_images
    from pathlib import Path

    fake_png = b"\x89PNG\r\n\x1a\nfake image data"
    buf = io.BytesIO()
    with _zf.ZipFile(buf, "w") as z:
        z.writestr("[Content_Types].xml",
                   '<?xml version="1.0"?><Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
                   '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
                   '<Default Extension="xml" ContentType="application/xml"/></Types>')
        z.writestr("_rels/.rels", '<?xml version="1.0"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"></Relationships>')
        z.writestr("word/media/image1.png", fake_png)
        z.writestr("word/media/unsupported.xyz", b"not an image")  # line 202: mime is None → skip
    docx_path = tmp_path / "con_imagen.docx"
    docx_path.write_bytes(buf.getvalue())

    images = _extract_docx_images(Path(docx_path))
    assert len(images) == 1  # only png, xyz is skipped
    assert images[0]["mime_type"] == "image/png"
    assert images[0]["data"] == fake_png


def test_load_doc_raises_value_error(tmp_path):
    doc = tmp_path / "old.doc"
    doc.write_bytes(b"data")
    with pytest.raises(ValueError, match=".doc.*Word 97"):
        dl.load_document(str(doc))


# ---------------------------------------------------------------------------
# _load_json — sin cambios
# ---------------------------------------------------------------------------

def test_load_json_returns_formatted_content(tmp_path):
    data = {"alumno": "Juan", "diagnostico": "TEA"}
    f = tmp_path / "paci.json"
    f.write_text(json.dumps(data), encoding="utf-8")

    result = dl.load_document(str(f), label="PACI JSON")

    assert "alumno" in result
    assert "Juan" in result
    assert "PACI JSON" in result


# ---------------------------------------------------------------------------
# Upload failure handling — Fix crítico PII
# ---------------------------------------------------------------------------

def test_load_pdf_upload_failure_propagates_cleanly(tmp_path):
    """Si files.upload falla, la excepción se propaga sin intentar delete."""
    pdf = tmp_path / "fail_upload.pdf"
    pdf.write_bytes(b"%PDF-1.4")

    client = MagicMock()
    client.files.upload.side_effect = RuntimeError("upload failed — quota exceeded")

    with patch("utils.document_loader.genai.Client", return_value=client), \
         patch.dict(os.environ, {"GOOGLE_API_KEY": "test-key"}):
        with pytest.raises(RuntimeError, match="upload failed"):
            dl.load_document(str(pdf))

    # delete NO debe llamarse si upload falló (uploaded es None)
    client.files.delete.assert_not_called()
